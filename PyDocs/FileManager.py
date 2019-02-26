# Handles all the file reading from .docx and regular files
from docx import Document

class WordFileManager:

    def __init__(self, path):
        self.document = path

    def setFilePath(self, path):
        self.document = path

    def readCells(self):
        wordDoc = Document(self.document)
        rows = []
        for table in wordDoc.tables:
            for row in table.rows:
                rows.append([cell.text for cell in row.cells])
        return rows

    def readText(self):
        wordDoc = Document(self.document)
        paras = []
        paras.append([p.text for p in wordDoc.paragraphs])
        return paras

class FileManager:

    def __init__(self, path):
        self.document = path

    def readFile(self):
        file = open(self.document, "r")
        return file.readlines()

if __name__ == "__main__":
    import os
    print(WordFileManager(os.getcwd() + "\PyDocs\demo.docx").readText())